"""Document parsers for PDF, DOCX, TEX, TXT, MD files, and arXiv HTML URLs."""

import base64
import os
import re
from pathlib import Path


def _tag_has_exact_class(tag, *target_classes: str) -> bool:
    """Check whether a BS4 tag has any of the given class names."""
    classes = tag.get("class", [])
    if isinstance(classes, str):
        classes = classes.split()
    return any(cls in target_classes for cls in classes)


def is_url(s: str) -> bool:
    """Check if a string looks like a URL."""
    return s.startswith("http://") or s.startswith("https://")


def _truncate_pdf_pages(path: Path, max_pages: int) -> Path:
    """Write a new PDF containing only the first max_pages pages. Returns the temp path."""
    import tempfile
    import pymupdf

    doc = pymupdf.open(str(path))
    if len(doc) <= max_pages:
        doc.close()
        return path
    print(f"  Truncating PDF from {len(doc)} to {max_pages} pages")
    doc.delete_pages(list(range(max_pages, len(doc))))
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    doc.save(tmp.name)
    doc.close()
    tmp.close()
    return Path(tmp.name)


def parse_document(
    file_path: str | Path,
    ocr: str | None = None,
    figures_dir: Path | None = None,
    max_pages: int | None = None,
) -> tuple[str, str, bool]:
    """Parse a document file or URL and return (title, full_text, was_ocr).

    Supported formats: .pdf, .docx, .tex, .txt, .md
    Also supports arXiv HTML URLs (e.g. https://arxiv.org/html/2310.06825).

    ocr: PDF OCR engine -- "mistral", "deepseek", "marker", "pymupdf", or None (auto).
    figures_dir: if provided, save extracted figures here (Mistral/DeepSeek OCR).
    was_ocr: True if the text went through OCR (PDF parsing), False otherwise.
    """
    path_str = str(file_path)

    if is_url(path_str):
        if "arxiv.org/abs/" in path_str:
            title, text, was_ocr = _parse_arxiv_abs(path_str, ocr=ocr)
            if was_ocr:
                from .ocr_postprocess import fix_ocr_notation
                text, corrections = fix_ocr_notation(text)
                for c in corrections:
                    print(f"  OCR fix: {c['old']} → {c['new']} ({c['reason']})")
            return title, text, was_ocr
        title, text = parse_arxiv_html(path_str)
        return title, text, False

    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        pdf_path = path
        if max_pages:
            pdf_path = _truncate_pdf_pages(path, max_pages)
        try:
            title, text, engine = _parse_pdf(pdf_path, ocr=ocr, figures_dir=figures_dir)
        finally:
            if pdf_path != path:
                pdf_path.unlink(missing_ok=True)
        # Post-process OCR output: fix notation errors
        from .ocr_postprocess import fix_ocr_notation
        text, corrections = fix_ocr_notation(text)
        for c in corrections:
            print(f"  OCR fix: {c['old']} → {c['new']} ({c['reason']})")
        # Store engine on module-level so callers (e.g. cmd_extract) can read it
        parse_document._last_ocr_engine = engine
        return title, text, True
    elif suffix == ".docx":
        title, text = _parse_docx(path)
        return title, text, False
    elif suffix == ".tex":
        title, text = _parse_tex(path)
        return title, text, False
    elif suffix in (".txt", ".md", ".markdown"):
        title, text, was_ocr = _parse_text(path)
        return title, text, was_ocr
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _parse_pdf(
    path: Path, ocr: str | None = None, figures_dir: Path | None = None
) -> tuple[str, str, str]:
    """Extract text from PDF.

    Engine priority (when ocr=None): Mistral OCR -> DeepSeek -> Marker -> pymupdf4llm.
    Set ocr="mistral", "deepseek", "marker", or "pymupdf" to force a specific engine.

    Returns (title, text, engine_used).
    """
    if ocr == "mistral":
        title, text = _parse_pdf_mistral(path, figures_dir=figures_dir)
        return title, text, "mistral"
    elif ocr == "deepseek":
        title, text = _parse_pdf_deepseek(path, figures_dir=figures_dir)
        return title, text, "deepseek"
    elif ocr == "marker":
        title, text = _parse_pdf_marker(path)
        return title, text, "marker"
    elif ocr == "pymupdf":
        title, text = _parse_pdf_pymupdf4llm(path)
        return title, text, "pymupdf"

    # Auto: try Mistral OCR first (best quality), then DeepSeek, Marker, pymupdf4llm
    if os.environ.get("MISTRAL_API_KEY"):
        try:
            title, text = _parse_pdf_mistral(path, figures_dir=figures_dir)
            return title, text, "mistral"
        except Exception as e:
            print(f"  Mistral OCR failed ({e}), trying next engine...")

    try:
        title, text = _parse_pdf_deepseek(path, figures_dir=figures_dir)
        return title, text, "deepseek"
    except (ImportError, ConnectionError, RuntimeError) as e:
        print(f"  DeepSeek OCR not available ({e}), trying Marker...")
    try:
        title, text = _parse_pdf_marker(path)
        return title, text, "marker"
    except (ImportError, FileNotFoundError, RuntimeError) as e:
        print(f"  Marker not available ({e}), trying pymupdf4llm...")

    title, text = _parse_pdf_pymupdf4llm(path)
    return title, text, "pymupdf4llm"


def _parse_pdf_mistral(path: Path, figures_dir: Path | None = None) -> tuple[str, str]:
    """High-quality PDF extraction using Mistral OCR via mistral-ocr-cli.

    Uses the OCRProcessor from mistral-ocr-cli, which provides:
    - Automatic retry with exponential backoff on transient errors
    - PDF chunking for documents >1000 pages
    - File upload API (avoids base64 size limits)
    - Figure extraction

    Cost: ~$0.001 per page.
    See also: https://github.com/r-uben/mistral-ocr-cli
    """
    from mistral_ocr import Config, OCRProcessor

    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        raise RuntimeError("MISTRAL_API_KEY not set")

    config = Config(
        api_key=api_key,
        model="mistral-ocr-latest",
        include_images=figures_dir is not None,
        table_format="markdown",
        extract_header=True,
        extract_footer=True,
        include_metadata=False,
        include_page_headings=False,
        quiet=True,
    )

    print(f"  Running Mistral OCR on {path.name}...")
    processor = OCRProcessor(config)
    result = processor.process_file(path)

    if not result or not result.get("success"):
        raise RuntimeError("Mistral OCR processing failed")

    response = result["response"]

    # Save images and build a mapping from OCR IDs to local paths
    image_map: dict[str, str] = {}
    n_images = 0
    if figures_dir is not None:
        figures_dir.mkdir(parents=True, exist_ok=True)
        for page in response.pages:
            images = getattr(page, "images", None) or []
            for idx, img in enumerate(images):
                b64_data = getattr(img, "image_base64", None)
                if not b64_data:
                    continue
                img_id = getattr(img, "id", None) or f"img-{idx}"
                ext = Path(img_id).suffix if "." in str(img_id) else ".png"
                filename = f"page{page.index + 1}_img{idx + 1}{ext}"
                save_path = figures_dir / filename
                save_path.write_bytes(base64.b64decode(b64_data))
                image_map[img_id] = f"figures/{filename}"
                n_images += 1

    # Concatenate page markdowns with table content
    pages = []
    for page in response.pages:
        page_parts = []
        if hasattr(page, "markdown") and page.markdown:
            page_parts.append(page.markdown)
        if hasattr(page, "tables") and page.tables:
            for table in page.tables:
                table_md = (
                    getattr(table, "content", None)
                    or getattr(table, "markdown", None)
                )
                if table_md:
                    page_parts.append(table_md)
        if page_parts:
            pages.append("\n\n".join(page_parts))

    if not pages:
        raise RuntimeError("Mistral OCR returned no content")

    markdown = "\n\n".join(pages)

    # Rewrite image references to point to saved files
    for img_id, local_path in image_map.items():
        markdown = markdown.replace(f"]({img_id})", f"]({local_path})")

    n_pages = len(response.pages)
    extras = f", {n_images} figures" if n_images else ""
    print(f"  Mistral OCR: {n_pages} pages extracted{extras} (~${n_pages * 0.001:.3f})")

    title = _extract_title_from_markdown(markdown)
    return title, markdown


def _parse_pdf_deepseek(path: Path, figures_dir: Path | None = None) -> tuple[str, str]:
    """PDF extraction using DeepSeek OCR via deepseek-ocr-cli (local).

    Runs a local vision model (DeepSeek-VL2) through Ollama or vLLM.
    Requires deepseek-ocr-cli installed and a running backend server.

    Install: pip install openaireview[deepseek]
    See also: https://github.com/r-uben/deepseek-ocr-cli
    """
    try:
        from deepseek_ocr import OCRProcessor as DeepSeekProcessor
    except ImportError:
        raise ImportError(
            "deepseek-ocr-cli not installed. "
            "Install with: pip install openaireview[deepseek]"
        )

    print(f"  Running DeepSeek OCR on {path.name}...")
    processor = DeepSeekProcessor(
        extract_images=figures_dir is not None,
        include_metadata=False,
        analyze_figures=figures_dir is not None,
    )

    # Ensure the backend model is loaded.
    # Uses private API — may break if deepseek-ocr-cli changes internals.
    if not processor._backend.model:
        processor._backend.load_model()

    result = processor.process_file(path, show_progress=True)
    markdown = result.output_text

    if not markdown.strip():
        raise RuntimeError("DeepSeek OCR returned no content")

    # Save figures if requested
    if figures_dir is not None:
        from deepseek_ocr.utils import sanitize_filename
        base_name = sanitize_filename(path.stem)
        src_figures = processor.output_dir / base_name / "figures"
        if src_figures.exists():
            import shutil
            figures_dir.mkdir(parents=True, exist_ok=True)
            for fig_file in src_figures.iterdir():
                shutil.copy2(fig_file, figures_dir / fig_file.name)

    n_pages = result.page_count
    print(f"  DeepSeek OCR: {n_pages} pages extracted in {result.processing_time:.1f}s")

    title = _extract_title_from_markdown(markdown)
    return title, markdown


def _parse_pdf_marker(path: Path) -> tuple[str, str]:
    """High-quality PDF extraction using Marker (preserves math as LaTeX).

    Tries the Python API first, then falls back to the Marker CLI.
    """
    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import text_from_rendered

        converter = PdfConverter(artifact_dict=create_model_dict())
        rendered = converter(str(path))
        markdown, _, _ = text_from_rendered(rendered)
    except ImportError:
        # Fall back to Marker CLI (avoids openai version conflict)
        import os
        import shutil
        import subprocess
        import tempfile

        marker_bin = shutil.which("marker_single")
        if not marker_bin:
            raise FileNotFoundError("marker_single not found on PATH")

        import platform
        env = os.environ.copy()
        if platform.system() == "Darwin":
            env["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
            env["TORCH_DEVICE"] = "cpu"

        with tempfile.TemporaryDirectory() as tmpdir:
            proc = subprocess.Popen(
                [marker_bin, str(path), "--output_dir", tmpdir],
                env=env,
            )
            try:
                proc.communicate(timeout=3600)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.communicate()
                raise RuntimeError("marker timed out after 3600s")
            if proc.returncode != 0:
                raise RuntimeError(f"marker failed (exit code {proc.returncode})")

            # Marker outputs to a subdirectory named after the PDF
            md_files = list(Path(tmpdir).rglob("*.md"))
            if not md_files:
                raise RuntimeError("marker produced no markdown output")
            markdown = md_files[0].read_text()

    title = _extract_title_from_markdown(markdown)
    return title, markdown


def _parse_pdf_pymupdf4llm(path: Path) -> tuple[str, str]:
    """PDF extraction using pymupdf4llm with GNN layout analysis.

    Fixes hyphenation, reading order, and table structure vs raw pymupdf.
    pymupdf-layout activates automatically when installed, enabling GNN-based
    table detection. Both packages are required dependencies.
    """
    try:
        import pymupdf.layout  # noqa: F401 — activates layout plugin
    except ImportError:
        pass  # layout plugin missing; pymupdf4llm still works, just without GNN

    import pymupdf4llm

    markdown = pymupdf4llm.to_markdown(str(path))
    markdown = _clean_pymupdf4llm_markdown(markdown)
    title = _extract_title_from_markdown(markdown)
    return title, markdown


def _clean_pymupdf4llm_markdown(md: str) -> str:
    """Post-process pymupdf4llm markdown for cleaner LLM input.

    - Strips the noisy '==> picture [WxH] intentionally omitted <==' lines.
      The pixel dimensions are meaningless and the phrasing distracts the LLM.
      Embedded figure text (chart labels, diagram text) and captions are kept.
    - Converts inline <br> separators in embedded text to newlines.
    """
    out = []
    for line in md.split("\n"):
        # Drop picture placeholder lines (with or without bold **)
        if "intentionally omitted" in line:
            stripped = line.strip().strip("*").strip()
            if stripped.startswith("==>"):
                continue
        # Clean up <br> in embedded figure text lines
        if "<br>" in line:
            line = line.replace("<br>", "\n")
        out.append(line)
    return "\n".join(out)


def _extract_title_from_markdown(markdown: str) -> str:
    """Extract the first heading from markdown text as the title."""
    fallback = ""
    for line in markdown.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            title = stripped.lstrip("# ").strip()
            # Strip bold markers that pymupdf4llm adds to headings
            return re.sub(r"\*\*(.+?)\*\*", r"\1", title)
        if not fallback:
            fallback = stripped[:200]
    return fallback


def _parse_docx(path: Path) -> tuple[str, str]:
    """Extract text from DOCX using python-docx."""
    import docx

    doc = docx.Document(str(path))
    paragraphs = []
    title = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if not title and para.style and para.style.name.startswith("Heading"):
            title = text

    full_text = "\n\n".join(paragraphs)

    if not title and paragraphs:
        title = paragraphs[0][:200]

    return title, full_text


def _parse_tex(path: Path) -> tuple[str, str]:
    """Extract text from LaTeX source."""
    text = path.read_text(encoding="utf-8", errors="replace")
    title = ""

    # Extract title from \title{...}
    title_match = re.search(r"\\title\{([^}]+)\}", text)
    if title_match:
        title = title_match.group(1).strip()
        # Clean common LaTeX artifacts from title
        title = re.sub(r"\\\\", " ", title)  # line breaks
        title = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", title)  # \textbf{X} → X
        title = re.sub(r"[{}]", "", title)  # stray braces
        title = re.sub(r"\s+", " ", title).strip()

    if not title:
        for line in text.split("\n"):
            if line.strip():
                title = line.strip()[:200]
                break

    return title, text


def _parse_text(path: Path) -> tuple[str, str, bool]:
    """Extract text from plain text or markdown.

    Detects YAML frontmatter with ocr_engine field (produced by `extract`).
    Returns (title, text_without_frontmatter, was_ocr).
    """
    raw = path.read_text(encoding="utf-8", errors="replace")
    was_ocr = False
    title = ""

    # Detect and parse YAML frontmatter
    text = raw
    if raw.startswith("---\n"):
        end = raw.find("\n---\n", 4)
        if end != -1:
            frontmatter = raw[4:end]
            text = raw[end + 5:]  # skip past closing ---\n
            # Parse frontmatter (simple key: value, no full YAML dep needed)
            for line in frontmatter.split("\n"):
                if line.startswith("title:"):
                    title = line.split(":", 1)[1].strip().strip('"')
                if line.startswith("ocr_engine:"):
                    was_ocr = True

    if not title:
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                title = stripped.lstrip("# ").strip()
            else:
                title = stripped[:200]
            break

    return title, text, was_ocr


def _parse_arxiv_abs(url: str, ocr: str | None = None) -> tuple[str, str, bool]:
    """Parse an arXiv abs URL: try HTML first, fall back to PDF."""
    html_url = re.sub(r"arxiv\.org/abs/", "arxiv.org/html/", url)
    try:
        title, text = parse_arxiv_html(html_url)
        return title, text, False
    except Exception as e:
        print(f"HTML version not available ({e}), falling back to PDF...")

    title, text = _fetch_arxiv_pdf(url, ocr=ocr)
    return title, text, True


def _fetch_arxiv_pdf(url: str, ocr: str | None = None) -> tuple[str, str]:
    """Fetch a PDF from an arXiv abs URL and parse it.

    Converts https://arxiv.org/abs/<id> to https://arxiv.org/pdf/<id>.
    """
    import tempfile
    from urllib.error import URLError
    from urllib.request import Request, urlopen

    pdf_url = re.sub(r"arxiv\.org/abs/", "arxiv.org/pdf/", url)
    print(f"Fetching PDF from {pdf_url}...")
    try:
        req = Request(pdf_url, headers={"User-Agent": "openaireview/0.1"})
        with urlopen(req, timeout=60) as resp:
            pdf_bytes = resp.read()
    except URLError as e:
        raise RuntimeError(f"Failed to fetch {pdf_url}: {e}") from e

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = Path(tmp.name)

    try:
        title, text, _engine = _parse_pdf(tmp_path, ocr=ocr)
        return title, text
    finally:
        tmp_path.unlink(missing_ok=True)


def _tabular_to_markdown(table_el) -> str:
    """Convert a BS4 ltx_tabular element to a markdown table."""
    rows = []
    for tr in table_el.find_all("tr"):
        cells = tr.find_all(["td", "th"])
        if not cells:
            continue
        row = [
            cell.get_text(" ", strip=True).replace("|", r"\|").replace("\n", " ")
            for cell in cells
        ]
        rows.append(row)

    if not rows:
        return ""

    ncols = max(len(r) for r in rows)

    def pad(r):
        return r + [""] * (ncols - len(r))

    lines = ["| " + " | ".join(pad(rows[0])) + " |",
             "| " + " | ".join(["---"] * ncols) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(pad(row)[:ncols]) + " |")
    return "\n".join(lines)


def _figure_or_table_to_markdown(fig_el) -> str:
    """Convert an ltx_figure or ltx_table element to markdown text.

    Tables become markdown tables (with caption above).
    Image figures are reduced to caption text only.
    """
    caption_el = fig_el.find(class_="ltx_caption")
    caption = (caption_el.get_text(" ", strip=True) if caption_el else "").replace("\n", " ")

    # Table figure: contains ltx_tabular
    tabular = fig_el.find(class_="ltx_tabular")
    if tabular:
        table_md = _tabular_to_markdown(tabular)
        if not table_md:
            return f"**{caption}**" if caption else ""
        return (f"**{caption}**\n\n{table_md}" if caption else table_md)

    # Image figure: collect main graphics, skip tiny caption icons
    imgs = [
        img for img in fig_el.find_all("img", class_="ltx_graphics")
        if not (caption_el and caption_el.find(lambda t: t is img))
        and int(img.get("width", "100") or "100") >= 30
    ]
    if not imgs:
        return f"**{caption}**" if caption else ""

    return f"*{caption}*" if caption else ""


def parse_arxiv_html(url: str) -> tuple[str, str]:
    """Fetch and parse an arXiv HTML page into (title, full_text).

    Works with arXiv HTML URLs like https://arxiv.org/html/2310.06825.
    The HTML is generated by LaTeXML and uses ltx_* CSS classes.
    Tables are converted to markdown tables; figures keep caption text only.
    """
    from urllib.error import URLError
    from urllib.request import Request, urlopen

    from bs4 import BeautifulSoup

    print(f"Fetching {url}...")
    try:
        req = Request(url, headers={"User-Agent": "openaireview/0.1"})
        with urlopen(req, timeout=30) as resp:
            html = resp.read().decode("utf-8", errors="replace")
    except URLError as e:
        raise RuntimeError(f"Failed to fetch {url}: {e}") from e

    soup = BeautifulSoup(html, "lxml")

    # Extract title
    title = ""
    title_el = soup.find(class_="ltx_title_document")
    if title_el:
        title = title_el.get_text(strip=True)
    if not title:
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)

    # Find the main document body
    doc = soup.find(class_="ltx_document") or soup.find("article") or soup.body
    if not doc:
        raise RuntimeError("Could not find paper content in HTML")

    # Remove bibliography, navigation, and other non-content elements
    for sel in ["nav", ".ltx_bibliography", ".ltx_TOC", "header", "footer",
                ".package-hierarchical-accordion", "#header", ".arxiv-watermark",
                ".ltx_role_affiliationtext"]:
        for el in doc.select(sel):
            el.decompose()

    # Pre-process figures and tables: convert to markdown and replace with
    # ltx_para marker divs so they appear at the correct position in the flow.
    # Use exact class match (not substring) to avoid matching ltx_figure_panel etc.
    inserted_markers = False
    for fig in doc.find_all(lambda tag: _tag_has_exact_class(tag, "ltx_figure", "ltx_table")):
        md = _figure_or_table_to_markdown(fig)
        if md:
            marker = soup.new_tag("div")
            marker["class"] = "ltx_para"
            marker["data-oar-content"] = md
            fig.replace_with(marker)
            inserted_markers = True
        else:
            fig.decompose()

    # Extract structured text using leaf content elements only.
    # ltx_para = paragraph text, ltx_title_* = headings, ltx_abstract = abstract,
    # ltx_theorem/ltx_proof = theorems. Captions are now handled via figure pre-processing.
    # We do NOT match ltx_section/ltx_subsection (containers that include all children).
    sections = []
    for element in doc.find_all(class_=re.compile(
        r"^ltx_(para$|title_|abstract$|theorem$|proof$)"
    )):
        # Figure/table markers: use pre-computed markdown directly
        oar_content = element.get("data-oar-content")
        if oar_content is not None:
            sections.append(oar_content)
            continue

        text = element.get_text(" ", strip=True)
        if not text:
            continue

        # Format headings
        cls = element.get("class", [])
        cls_str = " ".join(cls) if isinstance(cls, list) else cls
        if "ltx_title_document" in cls_str:
            sections.append(f"# {text}")
        elif "ltx_title_section" in cls_str:
            sections.append(f"\n## {text}")
        elif "ltx_title_subsection" in cls_str:
            sections.append(f"\n### {text}")
        elif "ltx_title_subsubsection" in cls_str:
            sections.append(f"\n#### {text}")
        elif "ltx_title_appendix" in cls_str:
            sections.append(f"\n## {text}")
        elif "ltx_title_abstract" in cls_str:
            # Skip — already handled by ltx_abstract match
            continue
        elif cls_str.startswith("ltx_title"):
            # Other titles (theorem, proof, etc.)
            sections.append(f"\n**{text}**")
        elif "ltx_abstract" in cls_str:
            # Extract just paragraph text, skip the title child
            abstract_paras = element.find_all(class_="ltx_p")
            if abstract_paras:
                abstract_text = "\n\n".join(
                    p.get_text(" ", strip=True) for p in abstract_paras
                )
            else:
                abstract_text = text
            sections.append(f"\n## Abstract\n{abstract_text}")
        else:
            sections.append(text)

    full_text = "\n\n".join(sections)

    # Fallback: if structured extraction got very little, use plain text
    if len(full_text) < 500 and not inserted_markers:
        full_text = doc.get_text("\n", strip=True)

    if not title:
        for line in full_text.split("\n"):
            if line.strip():
                title = line.strip()[:200]
                break

    return title, full_text
